import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.prompts import PromptTemplate
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import OpenAIEmbeddings
from langchain_cohere import CohereEmbeddings
from langchain_voyageai import VoyageAIEmbeddings
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_nvidia_ai_endpoints import NVIDIAEmbeddings
import os
import re
import time
from datetime import datetime
import requests
from bs4 import BeautifulSoup

# Disable file watcher to prevent inotify issues
os.environ['STREAMLIT_SERVER_FILE_WATCHER_TYPE'] = 'none'

# Model Configurations
AVAILABLE_MODELS = {
    "Gemma-2-9b": {
        "name": "gemma2-9b-it",
        "context_length": 1024,
        "description": "Gemma 2 9B IT model from Google",
        "temperature_range": (0.0, 1.0),
        "default_temperature": 0.3
    },
    "DeepSeek-R1-70b": {
        "name": "deepseek-r1-distill-llama-70b",
        "context_length": 4096,
        "description": "DeepSeek 70B distilled model",
        "temperature_range": (0.0, 1.0),
        "default_temperature": 0.3
    },
    "Llama-Scout-4-17B-Instruct": {
        "name": "meta-llama/llama-4-scout-17b-16e-instruct",
        "context_length": 1024,
        "description": "Meta Llama Scout 4 17B Instruct Model",
        "temperature_range": (0.0, 1.0),
        "default_temperature": 0.3
    },
    "qwen-qwq-32b": {
        "name": "qwen-qwq-32b",
        "context_length": 4096,
        "description": "Qwen QWQ 32B model from Alibaba Cloud",
        "temperature_range": (0.0, 1.0),
        "default_temperature": 0.3
    }
}

# Embeddings Configurations
AVAILABLE_EMBEDDINGS = {
    "BGE-small-en-v1.5": {
        "name": "BAAI/bge-small-en-v1.5",
        "description": "Lightweight, efficient embeddings model",
        "provider": "HuggingFace"
    },
    "OpenAI-text-embedding": {
        "name": "text-embedding-3-small",
        "description": "OpenAI's text embedding model",
        "provider": "OpenAI"
    },
    "Cohere-embed-english-v3.0": {
        "name": "embed-english-v3.0",
        "description": "Cohere embed-english-v3.0 model - high performance for semantic search.",
        "provider": "Cohere"
    },
    "Voyage-3-lite": {
        "name": "voyage-3-lite",
        "description": "VoyageAI embedding model optimized for latency and cost.",
        "provider": "VoyageAI"
    },
    "Google-text-embedding-004": {
        "name": "models/text-embedding-004",
        "description": "Google text embedding model for strong retrieval performance.",
        "provider": "Google"
    },
    "Nvidia-NV-Embed-v2": {
        "name": "NV-Embed-QA",
        "description": "Nvidia text embedding model.",
        "provider": "Nvidia"
    }
}

# Initialize session states
if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False

if 'url_input' not in st.session_state:
    st.session_state.url_input = ""

if 'config' not in st.session_state:
    st.session_state.config = {
        'k_value': 4,
        'chunk_size': 1000,
        'temperature': AVAILABLE_MODELS["DeepSeek-R1-70b"]["default_temperature"],
        'chunk_overlap': 200,
        'selected_model': "DeepSeek-R1-70b",
        'selected_embeddings': "BGE-small-en-v1.5"
    }  # Added missing closing brace

if 'model_metrics' not in st.session_state:
    st.session_state.model_metrics = {
        model: {
            'uses': 0,
            'avg_response_time': 0,
            'last_used': None
        } for model in AVAILABLE_MODELS
    }

if 'prev_model' not in st.session_state:
    st.session_state.prev_model = "DeepSeek-R1-70b"

def get_embeddings_model(embedding_choice):
    """Returns the appropriate embeddings model based on selection."""
    embedding_config = AVAILABLE_EMBEDDINGS[embedding_choice]
    try:
        if embedding_config["provider"] == "HuggingFace":
            return HuggingFaceEmbeddings(
                model_name=embedding_config["name"],
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        elif embedding_config["provider"] == "OpenAI":
            return OpenAIEmbeddings(
                model=embedding_config["name"],
                openai_api_key=st.secrets["OPENAI_API_KEY"]
            )
        elif embedding_config["provider"] == "Cohere": # New condition
            return CohereEmbeddings(
                model=embedding_config["name"],
                cohere_api_key=st.secrets["COHERE_API_KEY"] # Ensure this secret is available
            )
        elif embedding_config["provider"] == "VoyageAI":
            return VoyageAIEmbeddings(
                model=embedding_config["name"],
                voyage_api_key=st.secrets["VOYAGE_API_KEY"]
            )
        elif embedding_config["provider"] == "Google":
            return GoogleGenerativeAIEmbeddings(
                model=embedding_config["name"],
                google_api_key=st.secrets["GOOGLE_API_KEY"]
            )
        elif embedding_config["provider"] == "Nvidia":
            return NVIDIAEmbeddings(
                model=embedding_config["name"],
                nvidia_api_key=st.secrets["NVIDIA_API_KEY"]
            )
        else:
            raise ValueError(f"Unsupported embeddings provider: {embedding_config['provider']}")
    except Exception as e:
        st.error(f"Error initializing embeddings model: {str(e)}")
        return None

def get_word_text(doc):
    """Extracts text from a Word document."""
    try:
        document = Document(doc)
        text = ""
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return None

def get_document_text(docs):
    """Extracts text from uploaded PDF and Word documents and a URL."""
    text = ""
    url = st.session_state.get('url_input', '')

    try:
        for doc in docs:
            file_extension = doc.name.lower().split('.')[-1]

            if file_extension == 'pdf':
                pdf_reader = PdfReader(doc)
                for page in pdf_reader.pages:
                    text += page.extract_text()
            elif file_extension in ['docx', 'doc']:
                text += get_word_text(doc)
            text += "\n\n"

        if url:
            st.session_state.url_processed_successfully = False # Initialize as False
            try:
                response = requests.get(url, timeout=15) # Added timeout
                response.raise_for_status()  # Raise an exception for bad status codes (4xx or 5xx)
                
                soup = BeautifulSoup(response.content, 'html.parser')
                
                url_text_parts = []
                # Extract text from common content tags
                content_tags = soup.find_all(['p', 'article', 'main', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'span', 'div'])
                for tag in content_tags:
                    url_text_parts.append(tag.get_text(separator=" ", strip=True))
                
                url_text = "\n".join(url_text_parts)

                if not url_text.strip(): # Fallback or if primary extraction yields minimal text
                    url_text = soup.get_text(separator=" ", strip=True)

                if not url_text.strip():
                    st.warning(f"The URL ({url}) was fetched and parsed, but no text content was found.")
                    # url_processed_successfully remains False
                else:
                    text += url_text + "\n\n"
                    st.session_state.url_processed_successfully = True # Set to True on successful text extraction

            except requests.exceptions.MissingSchema:
                st.error(f"Invalid URL: '{url}'. Please include http:// or https:// at the beginning.")
            except requests.exceptions.Timeout:
                st.error(f"The request to the URL '{url}' timed out. Please try again later.")
            except requests.exceptions.HTTPError as e:
                if e.response.status_code == 404:
                    st.error(f"URL not found (404 Error): {url}")
                elif e.response.status_code == 403:
                    st.error(f"Access denied (403 Forbidden) for URL: {url}")
                elif e.response.status_code == 401:
                    st.error(f"Authentication required (401 Unauthorized) for URL: {url}")
                elif 400 <= e.response.status_code < 500:
                    st.error(f"Client error ({e.response.status_code}) accessing URL: {url}")
                elif 500 <= e.response.status_code < 600:
                    st.error(f"Server error ({e.response.status_code}) at URL: {url}. Please try again later.")
                else:
                    st.error(f"HTTP error ({e.response.status_code}) for URL: {url}. Details: {str(e)}")
            except requests.exceptions.ConnectionError:
                st.error(f"Failed to connect to the URL: '{url}'. Please check the URL and your internet connection.")
            except requests.exceptions.RequestException as e: # General fallback for other requests issues
                st.error(f"Error fetching URL '{url}': {str(e)}")
            except Exception as e: # Fallback for BeautifulSoup parsing or other unexpected errors
                st.error(f"Error parsing content from URL '{url}': {str(e)}")
                st.warning("Could not extract meaningful text content from the URL.")

        return text
    except Exception as e: # This is for errors related to file processing, keep it general
        st.error(f"Error reading document(s): {str(e)}")
        return None

def validate_selected_model():
    """Ensure selected model exists in AVAILABLE_MODELS."""
    if st.session_state.config['selected_model'] not in AVAILABLE_MODELS:
        st.session_state.config['selected_model'] = "DeepSeek-R1-70b"

def process_model_response(model_name, response_text):
    """Process model response based on model name."""
    if model_name.startswith(("deepseek", "qwen")):
        return re.sub(r'<think>.*?</think>', '', response_text, flags=re.DOTALL).strip()
    return response_text
    
def update_model_metrics(model_name, response_time):
    """Update usage metrics for the selected model."""
    metrics = st.session_state.model_metrics[model_name]
    metrics['uses'] += 1
    metrics['avg_response_time'] = (
        (metrics['avg_response_time'] * (metrics['uses'] - 1) + response_time) / metrics['uses']
    )
    metrics['last_used'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def get_text_chunks(text):
    """Splits extracted text into manageable chunks."""
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=st.session_state.config['chunk_size'],
            chunk_overlap=st.session_state.config['chunk_overlap'],
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = text_splitter.split_text(text)
        return chunks
    except Exception as e:
        st.error(f"Error splitting text: {str(e)}")
        return None

def get_vector_store(text_chunks):
    """Creates and saves a FAISS vector store from text chunks."""
    try:
        embeddings = get_embeddings_model(st.session_state.config['selected_embeddings'])
        if embeddings is None:
            return False

        vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
        vector_store.save_local("faiss_index")
        return True
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")
        return False

def get_conversational_chain():
    """Sets up a conversational chain using selected Groq model."""
    try:
        prompt_template = """
        You are a specialized assistant that ONLY answers questions based on the provided context.
        If the question cannot be answered using the information in the context, respond with:
        "I cannot answer this question as it's not covered in the provided documents."

        DO NOT use any external knowledge or make assumptions beyond what's in the context.
        If you're unsure about any part of the answer, err on the side of saying the information is not available.
        Keep the font style uniform in the response.

        Context:
        {context}

        Question:
        {question}

        Answer (based STRICTLY on the above context):
        """

        selected_model = st.session_state.config['selected_model']
        model_config = AVAILABLE_MODELS[selected_model]

        model = ChatGroq(
            temperature=st.session_state.config['temperature'],
            model_name=model_config['name'],
            groq_api_key=st.secrets["GROQ_API_KEY"]
        )
        prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
        chain = create_stuff_documents_chain(model, prompt)
        return chain
    except Exception as e:
        st.error(f"Error initializing Groq model: {str(e)}")
        return None

def compare_models(question, docs):
    """Compare responses from different models."""
    results = {}

    with st.expander("Model Comparison Results", expanded=True):
        st.markdown("### Model Comparison")
        st.markdown("Comparing responses from different models...")

        for model_name, model_config in AVAILABLE_MODELS.items():
            try:
                with st.spinner(f"Getting response from {model_name}..."):
                    start_time = time.time()

                    model = ChatGroq(
                        temperature=st.session_state.config['temperature'],
                        model_name=model_config['name'],
                        groq_api_key=st.secrets["GROQ_API_KEY"]
                    )

                    prompt_template = """
                    You are a specialized assistant that ONLY answers questions based on the provided context.
                    If the question cannot be answered using the information in the context, respond with:
                    "I cannot answer this question as it's not covered in the provided documents."

                    Context:
                    {context}

                    Question:
                    {question}

                    Answer (based STRICTLY on the above context):
                    """

                    prompt = PromptTemplate(
                        template=prompt_template,
                        input_variables=["context", "question"]
                    )
                    chain = create_stuff_documents_chain(model, prompt)

                    response = chain.invoke(
                        {"context": docs, "question": question}
                        # return_only_outputs is not a standard parameter for invoke with this chain type.
                        # The response itself is the string.
                    )

                    end_time = time.time()
                    response_time = end_time - start_time

                    response_text = process_model_response(model_name, response)

                    results[model_name] = {
                        'response': response_text,
                        'time': response_time
                    }

                    update_model_metrics(model_name, response_time)

                    st.markdown(f"**{model_name}**")
                    st.markdown(f"Response time: {response_time:.2f}s")
                    st.markdown(response_text)
                    st.markdown("---")
            except Exception as e:
                st.error(f"Error with {model_name}: {str(e)}")

    return results

def user_input(user_question):
    """Handles user queries by retrieving answers from the vector store."""
    try:
        embeddings = get_embeddings_model(st.session_state.config['selected_embeddings'])
        new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)

        docs = new_db.similarity_search(user_question, k=st.session_state.config['k_value'])

        if not docs:
            st.markdown("### Reply:\nI cannot answer this question as it's not covered in the provided documents.")
            return

        with st.expander("View source chunks used", expanded=False):
            st.markdown(f"**Using {st.session_state.config['k_value']} most relevant chunks:**")
            for i, doc in enumerate(docs, 1):
                st.markdown(f"**Chunk {i}:**\n{doc.page_content}\n---")

        if getattr(st.session_state, 'show_comparison', False):
            compare_models(user_question, docs)
            st.session_state.show_comparison = False
        else:
            start_time = time.time()
            chain = get_conversational_chain()
            if chain is None:
                return

            response = chain.invoke(
                {"context": docs, "question": user_question}
                # config={"return_only_outputs": True} # Not applicable here, response is the string.
            )

            end_time = time.time()
            response_time = end_time - start_time

            response_text = process_model_response(
                st.session_state.config['selected_model'],
                response
            )

            update_model_metrics(st.session_state.config['selected_model'], response_time)

            st.markdown(f"### Reply:\n{response_text}")
            st.markdown(f"*Response time: {response_time:.2f}s*")

    except Exception as e:
        st.error(f"Error processing question: {str(e)}")

def reset_app():
    """Resets the application state."""
    if os.path.exists("faiss_index"):
        try:
            import shutil
            shutil.rmtree("faiss_index")
        except Exception as e:
            st.error(f"Error clearing index: {str(e)}")
    st.session_state.processing_complete = False
    st.session_state.config = {
        'k_value': 4,
        'chunk_size': 1000,
        'temperature': AVAILABLE_MODELS["DeepSeek-R1-70b"]["default_temperature"],
        'chunk_overlap': 200,
        'selected_model': "DeepSeek-R1-70b",
        'selected_embeddings': "BGE-small-en-v1.5"
    }
    st.session_state.prev_model = "DeepSeek-R1-70b"
    st.rerun()

def main():
    """Main function to run the Streamlit app."""
    st.set_page_config(
        page_title="AI Powered Document Chat",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    st.title("AI Powered Document-Grounded Chat Assistant")

    # Sidebar
    with st.sidebar:
        validate_selected_model()

        st.header("Model Selection")
        selected_model = st.selectbox(
            "Choose Model",
            options=list(AVAILABLE_MODELS.keys()),
            index=list(AVAILABLE_MODELS.keys()).index(st.session_state.config['selected_model']),
            help="Select the Groq model to use for responses",
            key='model_selector'
        )

        if st.session_state.prev_model != selected_model:
            st.session_state.config['temperature'] = AVAILABLE_MODELS[selected_model]['default_temperature']
            st.session_state.prev_model = selected_model
            st.session_state.config['selected_model'] = selected_model

        with st.expander("Model Information", expanded=False):
            model_info = AVAILABLE_MODELS[selected_model]
            st.markdown(f"""
            **Model:** {selected_model}
            - **Description:** {model_info['description']}
            - **Context Length:** {model_info['context_length']} tokens
            - **Temperature Range:** {model_info['temperature_range'][0]} to {model_info['temperature_range'][1]}
            """)

        st.header("Embeddings Selection")
        selected_embeddings = st.selectbox(
            "Choose Embeddings Model",
            options=list(AVAILABLE_EMBEDDINGS.keys()),
            index=list(AVAILABLE_EMBEDDINGS.keys()).index(st.session_state.config['selected_embeddings']),
            help="Select the embeddings model to use for document processing"
        )

        # Update config with selected embeddings
        st.session_state.config['selected_embeddings'] = selected_embeddings

        with st.expander("Embeddings Information", expanded=False):
            embeddings_info = AVAILABLE_EMBEDDINGS[selected_embeddings]
            st.markdown(f"""
            **Model:** {selected_embeddings}
            - **Description:** {embeddings_info['description']}
            - **Provider:** {embeddings_info['provider']}
            """)

        st.session_state.config['selected_embeddings'] = selected_embeddings

        st.subheader("Model Parameters")
        temperature = st.slider(
            "Temperature",
            min_value=AVAILABLE_MODELS[selected_model]['temperature_range'][0],
            max_value=AVAILABLE_MODELS[selected_model]['temperature_range'][1],
            value=st.session_state.config['temperature'],
            step=0.1,
            help="Controls randomness in the response. Lower values make responses more focused."
        )

        st.subheader("Chunk Parameters")
        col1, col2 = st.columns(2)
        with col1:
            chunk_size = st.number_input(
                "Chunk Size",
                min_value=100,
                max_value=2000,
                value=st.session_state.config['chunk_size'],
                step=100,
                help="Size of text chunks during processing"
            )

        with col2:
            chunk_overlap = st.number_input(
                "Chunk Overlap",
                min_value=0,
                max_value=500,
                value=st.session_state.config['chunk_overlap'],
                step=50,
                help="Overlap between chunks"
            )

        st.subheader("Search Parameters")
        k_value = st.slider(
            "Number of chunks (k)",
            min_value=1,
            max_value=10,
            value=st.session_state.config['k_value'],
            help="Number of relevant chunks to use for answering"
        )

        st.subheader("Model Comparison")
        if st.button("Compare Models"):
            st.session_state.show_comparison = True

        st.session_state.config.update({
            'temperature': temperature,
            'chunk_size': chunk_size,
            'chunk_overlap': chunk_overlap,
            'k_value': k_value,
            'selected_model': selected_model
        })

        if st.checkbox("Show Model Metrics"):
            st.markdown("### Model Usage Metrics")
            for model, metrics in st.session_state.model_metrics.items():
                st.markdown(f"""
                **{model}**
                - Uses: {metrics['uses']}
                - Avg Response Time: {metrics['avg_response_time']:.2f}s
                - Last Used: {metrics['last_used'] or 'Never'}
                """)

        st.markdown("---")

        st.header("Document Upload")
        uploaded_docs = st.file_uploader(
            "Upload documents:",
            accept_multiple_files=True,
            type=["pdf", "docx", "doc"],
            key="file_uploader"
        )
        st.session_state.url_input = st.text_input("Or enter URL to ingest content:", placeholder="https://example.com", key="url_input_widget")

        processed_items = []
        if uploaded_docs:
            for doc in uploaded_docs:
                processed_items.append(doc.name)
        if st.session_state.get('url_input', ''):
            processed_items.append(f"URL: {st.session_state.url_input}")

        if not uploaded_docs and not st.session_state.get('url_input', ''):
            st.info("Please upload PDF or Word documents, or enter a URL to begin.")
            st.button("Process Documents", disabled=True, key="process_disabled")
        else:
            if processed_items:
                st.markdown("### Items to Process:")
                for item in processed_items:
                    st.markdown(f"- {item}")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("Process Documents", type="primary", key="process"):
                    with st.spinner("Processing items..."):
                        # Reset flag for URL processing success message
                        st.session_state.url_processed_successfully = False 
                        raw_text = get_document_text(uploaded_docs) # uploaded_docs can be empty
                        if raw_text:
                            text_chunks = get_text_chunks(raw_text)
                            if text_chunks and get_vector_store(text_chunks):
                                st.session_state.processing_complete = True
                                success_messages = []
                                if uploaded_docs:
                                    success_messages.append(f"{len(uploaded_docs)} document(s) processed.")
                                if st.session_state.get('url_input', '') and st.session_state.get('url_processed_successfully', False):
                                    success_messages.append("URL content processed.")
                                elif st.session_state.get('url_input', '') and not st.session_state.get('url_processed_successfully', False):
                                     # Error message for URL processing failure is handled in get_document_text
                                     pass # No specific success message if URL processing failed
                                
                                if success_messages:
                                    st.success(" ".join(success_messages))
                                else:
                                    # This case might happen if URL fails and no files were uploaded
                                    st.warning("No content was successfully processed.") 
                                st.rerun()
                            elif not text_chunks and st.session_state.get('url_input', '') and not uploaded_docs:
                                # Handles case where only URL was provided but text extraction failed or was empty
                                st.error("Failed to extract text from the URL or the URL content was empty. Please check the URL and try again.")
                        elif not uploaded_docs and st.session_state.get('url_input', ''): 
                            # This case handles if get_document_text returns None AND only URL was given
                            # Error for URL fetching itself is in get_document_text
                            st.error("Failed to process the URL. Please ensure it's valid and accessible.")
                        elif not raw_text and uploaded_docs:
                             st.error("Failed to extract text from uploaded documents.")

            with col2:
                if st.button("Clear All", type="secondary", key="clear"):
                    reset_app()

    st.markdown("### Ask Questions")
    st.markdown("The assistant will only answer questions based on the uploaded documents.")

    if not st.session_state.processing_complete:
        st.text_input(
            "Your question:",
            placeholder="Process documents first...",
            disabled=True
        )
        st.info("Please upload and process documents to start asking questions.")
    else:
        user_question = st.text_input(
            "Your question:",
            placeholder="Ask about the uploaded documents..."
        )
        if user_question:
            with st.spinner("Finding answer..."):
                user_input(user_question)

    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center'>
            <p style='color: #666; font-size: 0.8em;'>
                AI Powered Document-Grounded Chat Assistant | Powered by Groq
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
